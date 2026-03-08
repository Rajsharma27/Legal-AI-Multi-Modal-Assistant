import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Union

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = {
    ".pdf": "pdf",
    ".txt": "text",
    ".docx": "docx",
    ".doc": "docx",
}


@dataclass
class Document:
    """A loaded document with its text and basic metadata."""
    file_path: str
    file_name: str
    file_type: str
    text: str
    page_count: int = 1
    file_size_bytes: int = 0
    loaded_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())


# ---------- Format-specific readers ----------

def _read_pdf(path: Path) -> tuple[str, int]:
    """Extract text from a PDF. Returns (text, page_count)."""
    from PyPDF2 import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            pages.append(page_text)
    return "\n\n".join(pages), len(reader.pages)


def _read_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    import docx

    doc = docx.Document(str(path))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def _read_text(path: Path) -> str:
    """Read a plain-text file."""
    return path.read_text(encoding="utf-8", errors="replace")


# ---------- Public API ----------

def load_document(path: Union[str, Path]) -> Document:
    """
    Load a single document and return its extracted text.

    Args:
        path: Path to a PDF, DOCX, or TXT file.

    Returns:
        A Document dataclass with text and metadata.

    Raises:
        FileNotFoundError: If the file doesn't exist.
        ValueError: If the format is unsupported.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    ext = path.suffix.lower()
    fmt = SUPPORTED_FORMATS.get(ext)
    if fmt is None:
        raise ValueError(f"Unsupported format: {ext}")

    logger.info("Loading %s file: %s", fmt, path.name)

    text = ""
    page_count = 1

    if fmt == "pdf":
        text, page_count = _read_pdf(path)
    elif fmt == "docx":
        text = _read_docx(path)
    elif fmt == "text":
        text = _read_text(path)

    return Document(
        file_path=str(path.resolve()),
        file_name=path.name,
        file_type=fmt,
        text=text.strip(),
        page_count=page_count,
        file_size_bytes=path.stat().st_size,
    )


def load_directory(
    directory: Union[str, Path],
    recursive: bool = True,
) -> list[Document]:
    """
    Load all supported documents from a directory.

    Args:
        directory: Folder to scan.
        recursive: Search subdirectories too.

    Returns:
        List of Document objects.
    """
    directory = Path(directory)
    if not directory.is_dir():
        raise NotADirectoryError(f"Not a directory: {directory}")

    pattern = "**/*" if recursive else "*"
    files = sorted(
        p for p in directory.glob(pattern)
        if p.suffix.lower() in SUPPORTED_FORMATS
    )
    logger.info("Found %d documents in %s", len(files), directory)

    docs = []
    for f in files:
        try:
            docs.append(load_document(f))
        except Exception as e:
            logger.error("Failed to load %s: %s", f.name, e)
    return docs


def get_supported_extensions() -> list[str]:
    """Return all file extensions the loader can handle."""
    return sorted(SUPPORTED_FORMATS.keys())